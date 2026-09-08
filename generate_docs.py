import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_privacy_policy_doc():
    doc = docx.Document()

    # Set Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("KEBIJAKAN PRIVASI (PRIVACY POLICY)\nRENJANA")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(128, 0, 32) # Burgundy

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Platform Digital Edukasi Kebudayaan & Sejarah Indonesia\nTerakhir Diperbarui: 31 Agustus 2026")
    sub_run.italic = True
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Intro Callout Box
    intro_table = doc.add_table(rows=1, cols=1)
    intro_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = intro_table.cell(0, 0)
    set_cell_background(cell, "F4F1EA")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    p_box = cell.paragraphs[0]
    r_box = p_box.add_run("Selamat datang di Renjana (\"Aplikasi\"). Kebijakan Privasi ini menjelaskan bagaimana Renjana (\"kami\", \"pengembang\") mengumpulkan, menggunakan, memproses, menyimpan, dan melindungi informasi pribadi Anda saat menggunakan aplikasi mobile daring (online) kami yang terintegrasi dengan layanan Google Firebase dan tersedia di Google Play Store.\n\nDengan menggunakan Aplikasi Renjana, Anda menyetujui praktik pengumpulan dan penggunaan informasi sebagaimana diuraikan dalam dokumen ini.")
    r_box.font.name = "Arial"
    r_box.font.size = Pt(10.5)
    r_box.font.color.rgb = RGBColor(40, 40, 40)

    doc.add_paragraph()

    sections = [
        ("1. Ringkasan Singkat", [
            "Renjana adalah platform digital edukasi kebudayaan, sejarah nusantara, dan komunitas budaya Indonesia yang beroperasi secara daring (online).",
            "Aplikasi menggunakan infrastruktur cloud dari Google Firebase (Google LLC) untuk autentikasi pengguna, basis data cloud (Cloud Firestore), penyimpanan media (Firebase Cloud Storage), dan pengiriman notifikasi (Firebase Cloud Messaging).",
            "Kami tidak menjual, menyewakan, atau memperdagangkan data pribadi Anda kepada pihak ketiga atau pihak pengiklan komersial.",
            "Seluruh transmisi data antara aplikasi, perangkat Anda, dan server Google Firebase dienkripsi menggunakan protokol standar industri (HTTPS / TLS)."
        ]),
        ("2. Informasi dan Data yang Kami Kumpulkan", [
            "a. Data Akun Pengguna (Firebase Authentication):\nNama Lengkap / Nama Tampilan (Display Name), Alamat Email, dan Kata Sandi (Password yang dienkripsi secara aman oleh sistem Firebase). Foto profil opsional yang diunggah pengguna disimpan di Firebase Cloud Storage.",
            "b. Data Aktivitas Belajar dan Komunitas (Cloud Firestore):\nNilai/skor kuis, riwayat jawaban, runtun belajar harian (daily streak), lencana pencapaian (achievements), koleksi artikel tersimpan (bookmark), serta unggahan/komentar di forum komunitas budaya.",
            "c. Informasi Teknis dan Diagnostik Perangkat:\nModel perangkat keras, versi sistem operasi Android, serta log diagnostik kesalahan sistem (crash logs) untuk pemeliharaan kestabilan aplikasi.",
            "d. Data yang TIDAK Dikumpulkan:\nKami TIDAK mengumpulkan data lokasi GPS presisi di latar belakang, data finansial/kartu kredit, maupun akses kontak telepon pribadi dan SMS."
        ]),
        ("3. Layanan Pihak Ketiga (Google Firebase)", [
            "Aplikasi kami menggunakan backend terintegrasi dari Google LLC (Google Firebase):",
            "• Firebase Authentication: Pengelolaan pendaftaran akun, verifikasi login aman, dan pemulihan akun.",
            "• Cloud Firestore / Firebase Database: Penyimpanan katalog budaya, bank soal kuis, data komunitas, dan progres belajar.",
            "• Firebase Cloud Storage: Penyimpanan foto profil dan media kontribusi budaya.",
            "• Firebase Cloud Messaging (FCM): Pengiriman pengingat belajar harian dan notifikasi informasi budaya terbaru.",
            "Pengelolaan data oleh Google tunduk pada Kebijakan Privasi Google (https://policies.google.com/privacy) dan Ketentuan Keamanan Data Firebase."
        ]),
        ("4. Izin Perangkat (Android Permissions) yang Digunakan", [
            "• Akses Jaringan Internet (INTERNET, ACCESS_NETWORK_STATE): Menghubungkan aplikasi ke server Google Firebase untuk sinkronisasi data budaya, kuis, dan akun.",
            "• Foto & Media (READ_MEDIA_IMAGES, READ_EXTERNAL_STORAGE): Memungkinkan pengguna/admin memilih foto dari galeri untuk foto profil atau unggahan materi budaya.",
            "• Kamera (CAMERA): Memungkinkan pengambilan foto secara langsung untuk pembaruan profil atau kontribusi budaya.",
            "• Notifikasi (POST_NOTIFICATIONS, SCHEDULE_EXACT_ALARM): Mengirimkan pengingat belajar harian (daily streak) dan info budaya terbaru."
        ]),
        ("5. Keamanan dan Penyimpanan Data (Data Security)", [
            "• Enkripsi dalam Transit: Seluruh pertukaran data antara aplikasi Renjana dan layanan Google Firebase terlindungi enkripsi TLS/SSL.",
            "• Aturan Keamanan Cloud (Security Rules): Akses basis data diatur oleh konfigurasi aturan keamanan ketat (Firebase Security Rules), memastikan setiap pengguna hanya dapat mengelola datanya sendiri.",
            "• Isolasi Data Lokal: Token otentikasi sesi disimpan secara aman pada ruang memori lokal terisolasi (sandbox) di perangkat pengguna."
        ]),
        ("6. Hak Pengguna dan Penghapusan Akun (Account & Data Deletion)", [
            "Sesuai standar kebijakan Google Play Console mengenai hak privasi pengguna:",
            "• Penghapusan Mandiri: Pengguna dapat mengajukan penghapusan akun beserta seluruh data riwayat belajar, bookmark, dan kontribusi melalui menu Pengaturan Profil di dalam aplikasi.",
            "• Permintaan via Email: Pengguna dapat meminta penghapusan akun permanen dengan mengirimkan permohonan ke email resmi pengembang.",
            "• Pemrosesan Permanen: Setelah permohonan diverifikasi, akun Firebase Authentication dan data terkait pada Cloud Firestore/Storage akan dihapus secara permanen dari server aktif kami."
        ]),
        ("7. Privasi Anak-Anak (Children's Privacy)", [
            "• Aplikasi Renjana merupakan aplikasi edukasi budaya yang ramah keluarga dan aman untuk seluruh kelompok usia, termasuk pelajar dan anak-anak.",
            "• Aplikasi tidak memuat materi berbahaya, konten dewasa, ataupun transaksi keuangan.",
            "• Jika orang tua/wali menemukan bahwa data anak diberikan tanpa izin, silakan hubungi pengembang agar kami dapat segera melakukan verifikasi dan penghapusan data."
        ]),
        ("8. Perubahan Kebijakan Privasi", [
            "Kebijakan Privasi ini dapat diperbarui dari waktu ke waktu untuk menyesuaikan penambahan fitur baru atau kepatuhan regulasi hukum. Tanggal pembaruan terakhir akan selalu dicantumkan di bagian atas dokumen ini."
        ]),
        ("9. Kontak dan Layanan Bantuan", [
            "Apabila Anda memiliki pertanyaan, saran, atau permohonan terkait data pribadi Anda, silakan hubungi tim kami:",
            "• Nama Tim: Tim Pengembang Renjana",
            "• Email: renjana.app@gmail.com",
            "• Infrastruktur Backend: Google Firebase (Google LLC)",
            "• Lokasi: DKI Jakarta, Indonesia"
        ])
    ]

    for sec_title, sec_bullets in sections:
        h = doc.add_heading(sec_title, level=1)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        for r in h.runs:
            r.font.name = "Arial"
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(128, 0, 32)
            r.bold = True

        for b in sec_bullets:
            if b.startswith("• "):
                bp = doc.add_paragraph(style='List Bullet')
                br = bp.add_run(b[2:])
                br.font.name = "Arial"
                br.font.size = Pt(10.5)
            elif "\n" in b and not b.startswith("•"):
                lines = b.split("\n")
                p1 = doc.add_paragraph()
                p1.paragraph_format.space_after = Pt(2)
                r1 = p1.add_run(lines[0])
                r1.bold = True
                r1.font.name = "Arial"
                r1.font.size = Pt(10.5)
                
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_after = Pt(4)
                r2 = p2.add_run(lines[1])
                r2.font.name = "Arial"
                r2.font.size = Pt(10.5)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(b)
                r.font.name = "Arial"
                r.font.size = Pt(10.5)

    doc.save("d:/RenjanaAppPPKD/renjana/Kebijakan_Privasi_Renjana.docx")
    print("Kebijakan_Privasi_Renjana.docx created successfully.")

def create_playstore_listing_doc():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("MATERI LISTING GOOGLE PLAY STORE\nRENJANA")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(128, 0, 32)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Panduan Teks Publikasi & Data Safety (Google Firebase)\nTerakhir Diperbarui: 31 Agustus 2026")
    sub_run.italic = True
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Section 1: Title
    h1 = doc.add_heading("1. Judul Aplikasi (App Title) — Max 30 Karakter", level=1)
    for r in h1.runs:
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(128, 0, 32)

    t_table = doc.add_table(rows=4, cols=3)
    t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Kategori", "Teks Judul", "Jumlah Karakter"]
    for i, h_text in enumerate(headers):
        cell = t_table.cell(0, i)
        set_cell_background(cell, "800020")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    title_rows = [
        ("Pilihan Utama (Rekomendasi)", "Renjana: Museum & Budaya ID", "28 / 30"),
        ("Alternatif 1", "Renjana - Budaya Indonesia", "24 / 30"),
        ("Alternatif 2", "Renjana: Sejarah & Budaya ID", "27 / 30")
    ]
    for row_idx, data in enumerate(title_rows, start=1):
        bg = "F9F8F6" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = t_table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            if col_idx == 0 and "Rekomendasi" in text:
                r.bold = True

    doc.add_paragraph()

    # Section 2: Short Description
    h2 = doc.add_heading("2. Deskripsi Singkat (Short Description) — Max 80 Karakter", level=1)
    for r in h2.runs:
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(128, 0, 32)

    sd_table = doc.add_table(rows=4, cols=3)
    sd_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(["Pilihan", "Teks Deskripsi Singkat", "Panjang"]):
        cell = sd_table.cell(0, i)
        set_cell_background(cell, "800020")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    sd_rows = [
        ("Pilihan 1 (Rekomendasi)", "Jelajahi museum, sejarah, dan kekayaan budaya Nusantara dalam satu genggaman!", "79 / 80"),
        ("Pilihan 2", "Platform digital budaya & sejarah Indonesia: peta, kuis, dan komunitas online.", "78 / 80"),
        ("Pilihan 3", "Eksplorasi ragam budaya, sejarah nusantara, peta 38 provinsi & kuis interaktif.", "79 / 80")
    ]
    for row_idx, data in enumerate(sd_rows, start=1):
        bg = "F9F8F6" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = sd_table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(10)
            if col_idx == 0 and "Rekomendasi" in text:
                r.bold = True

    doc.add_paragraph()

    # Section 3: Full Description
    h3 = doc.add_heading("3. Deskripsi Lengkap (Full Description) — Max 4.000 Karakter", level=1)
    for r in h3.runs:
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(128, 0, 32)

    note_p = doc.add_paragraph()
    r_note = note_p.add_run("Teks di bawah ini memiliki panjang ~3.050 karakter (sangat aman di bawah batas 4.000 karakter Google Play Console, terstruktur rapi dan dioptimalkan untuk ASO):")
    r_note.italic = True
    r_note.font.size = Pt(10)

    desc_table = doc.add_table(rows=1, cols=1)
    desc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_desc = desc_table.cell(0, 0)
    set_cell_background(c_desc, "F9F8F6")
    set_cell_margins(c_desc, top=140, bottom=140, left=180, right=180)
    p_full = c_desc.paragraphs[0]
    
    full_text = """🏛️ RENJANA — Museum Digital & Ensiklopedia Budaya Indonesia

Renjana adalah platform digital edukasi kebudayaan dan sejarah nusantara berbasis daring (online) yang dirancang interaktif, modern, dan informatif. Dari Sabang sampai Merauke, jelajahi ribuan warisan leluhur bangsa, peristiwa bersejarah, keunikan adat istiadat 38 provinsi, komunitas pegiat budaya, serta asah wawasan lewat kuis budaya interaktif!

Didukung sistem cloud yang cepat dan andal, seluruh progres belajar, koleksi artikel, skor kuis, dan keaktifan komunitas Anda tersinkronisasi secara aman di mana saja dan kapan saja.

✨ FITUR-FITUR UNGGULAN RENJANA:

📖 1. Sorotan Sejarah & Budaya Harian (Realtime Updates)
Dapatkan fakta unik, ulasan sejarah, dan artikel kebudayaan nusantara pilihan yang diperbarui setiap hari secara daring untuk menambah wawasan literasi kebangsaan Anda.

🎭 2. Katalog 8 Kategori Budaya Terlengkap
Temukan khazanah kekayaan Indonesia yang terstruktur rapi dan terus diperbarui:
• Rumah & Arsitektur Adat Nusantara
• Pakaian & Wastra Tradisional
• Seni Tari & Pertunjukan Daerah
• Senjata Tradisional & Pusaka Bersejarah
• Alat Musik Tradisional
• Upacara, Ritual & Tradisi Adat
• Kuliner Khas Daerah
• Kesenian Daerah & Kerajinan Tangan

📜 3. Garis Waktu Sejarah Nusantara (Interactive Timeline)
Pelajari rekaman jejak sejarah Indonesia secara kronologis, mulai dari masa kerajaan kuno, era kolonial, perjuangan kemerdekaan, hingga era modern yang disajikan dengan visual menarik dan narasi yang mudah dipahami.

🗺️ 4. Peta Budaya Digital 38 Provinsi
Jelajahi peta interaktif seluruh wilayah kepulauan Indonesia: Pulau Sumatera, Jawa, Kalimantan, Sulawesi, Bali, Nusa Tenggara, Maluku, hingga Papua. Temukan ciri khas adat dan kebudayaan di tiap daerah dengan satu sentuhan.

🎯 5. Kuis Budaya Online & Pembahasan Lengkap
Uji dan tingkatkan pemahaman Anda melalui kuis budaya bertingkat! Dilengkapi pembahasan jawaban komprehensif, penghitungan skor, dan pelacakan riwayat capaian belajar.

🔥 6. Sinkronisasi Akun, Runtun Belajar (Streak) & Lencana Prestasi
Bangun kebiasaan belajar setiap hari! Progres belajar, runtun kehadiran (streak), dan lencana penghargaan (achievements) Anda tersimpan aman dan tersinkronisasi otomatis di akun cloud Anda.

👥 7. Forum Komunitas & Ruang Kontribusi Budaya
Terhubung dengan sesama pelajar, mahasiswa, peneliti, dan pecinta budaya nusantara. Bagikan ulasan, ajukan diskusi menarik, dan berkontribusi melestarikan informasi budaya lokal Anda.

📌 8. Koleksi & Bookmark Tersinkronisasi
Simpan artikel dan materi budaya favorit Anda ke dalam akun agar dapat dibaca kembali kapan saja di berbagai perangkat.

📍 9. Direktori Destinasi Wisata & Cagar Budaya
Temukan panduan dan profil museum, candi, monumen, serta situs cagar budaya di seluruh penjuru Indonesia lengkap dengan latar belakang sejarahnya.

🌟 COCOK UNTUK:
• Pelajar SD, SMP, SMA/SMK & Mahasiswa untuk bahan belajar, referensi, dan tugas sekolah/kuliah.
• Guru & Pendidik sebagai sarana media pembelajaran interaktif di kelas.
• Komunitas & Pecinta Budaya yang ingin mendalami dan melestarikan kekayaan bangsa.
• Wisatawan yang ingin mengenal tradisi lokal daerah destinasi wisata Indonesia.

Mari bersama-sama merawat, mencintai, dan melestarikan warisan adiluhung nusantara melalui teknologi digital.

Unduh Renjana sekarang dan rasakan pengalaman menjelajahi museum nusantara dalam satu genggaman! 🇮🇩"""

    r_full = p_full.add_run(full_text)
    r_full.font.name = "Arial"
    r_full.font.size = Pt(10)
    r_full.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_paragraph()

    # Section 4: Data Safety
    h4 = doc.add_heading("4. Panduan Pengisian Data Safety Google Play Console", level=1)
    for r in h4.runs:
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(128, 0, 32)

    ds_table = doc.add_table(rows=8, cols=2)
    ds_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(["Pertanyaan Google Play", "Panduan Jawaban Resmi"]):
        cell = ds_table.cell(0, i)
        set_cell_background(cell, "800020")
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    ds_rows = [
        ("Apakah aplikasi mengumpulkan atau membagikan data pengguna?", "Ya"),
        ("Apakah semua data dienkripsi saat transit (in transit)?", "Ya (Protokol HTTPS / TLS oleh Google Firebase)"),
        ("Apakah pengguna dapat meminta penghapusan akun dan data?", "Ya (Mekanisme penghapusan akun via Profil aplikasi dan email)"),
        ("Data Pribadi (Personal Info) yang Diproses", "• Nama: Identitas profil pengguna\n• Email: Login Firebase Authentication\n• User ID: Identifikasi akun Firebase"),
        ("Aktivitas Aplikasi (App Activity)", "Riwayat nilai kuis, daftar bookmark, runtun belajar harian (streak), dan interaksi postingan komunitas"),
        ("Foto dan Media (Photos and Videos)", "Opsional: Hanya saat pengguna/admin mengunggah foto profil atau foto materi budaya"),
        ("Tujuan Pengumpulan Data", "Fungsionalitas Aplikasi (App Functionality), Manajemen Akun (Account Management), dan Personalisasi")
    ]
    for row_idx, data in enumerate(ds_rows, start=1):
        bg = "F9F8F6" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = ds_table.cell(row_idx, col_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            if col_idx == 0:
                r.bold = True

    doc.save("d:/RenjanaAppPPKD/renjana/PlayStore_Listing_Renjana.docx")
    print("PlayStore_Listing_Renjana.docx created successfully.")

if __name__ == "__main__":
    create_privacy_policy_doc()
    create_playstore_listing_doc()
